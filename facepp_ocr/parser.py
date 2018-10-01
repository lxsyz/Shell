#!/usr/bin/env python
# -*- coding: UTF-8 -*-
from bs4 import BeautifulSoup
import os
import re

from docx import Document
from docx.oxml.ns import qn
from docx.shared import RGBColor

colorlist = ["gray", "green", "blue", "red", "deeppink"]

def html2word(document, res_list, font_list, num_list, lang, filename):
    rows = 0
    for content in res_list:
        p = document.add_paragraph()

        for index in range(len(content)):
            if lang == "cn":
                character = content[index:index + 1]
                if character == ";":
                    character = u"；"
                elif character == "[":
                    character = u"【"
                elif character == "]":
                    character = u"】"
                elif character == "(":
                    character = u"（"
                elif character == ")":
                    character = u"）"

            colored = False
            for (row, col, color) in font_list:
                if row > rows:
                    break
                if col == index and row == rows:
                    font = p.add_run(character).font
                    if color == 'gray':
                        font.color.rgb = RGBColor(128, 128, 128)
                        num_list[0] += 1
                    elif color == 'green':
                        font.color.rgb = RGBColor(0, 128, 0)
                        num_list[1] += 1
                    elif color == 'blue':
                        font.color.rgb = RGBColor(0, 0, 225)
                        num_list[2] += 1
                    elif color == 'red':
                        font.color.rgb = RGBColor(255, 0, 0)
                        num_list[3] += 1
                    # elif color == 'deeppink':
                    #     font.bold=True
                    #     font.italic=True
                    colored = True
                    break
            if not colored:
                p.add_run(character)
        rows += 1
    # p = document.add_paragraph()
    # p.add_run(filename)

def html2word2(document, res_list, font_list, num_list, lang, filename):
    rows = 0
    for content in res_list:
        p = document.add_paragraph()

        for index in range(len(content)):
            if lang == "cn":
                character = content[index:index + 1]
                if character == ";":
                    character = u"；"
                elif character == "[":
                    character = u"【"
                elif character == "]":
                    character = u"】"
                elif character == "(":
                    character = u"（"
                elif character == ")":
                    character = u"）"

            colored = False
            for (row, col, color) in font_list:
                if row > rows:
                    break
                if col == index and row == rows:
                    font = p.add_run(character).font
                    if color == 'gray':
                        font.color.rgb = RGBColor(128, 128, 128)
                        num_list[0] += 1
                    elif color == 'green':
                        font.color.rgb = RGBColor(0, 128, 0)
                        num_list[1] += 1
                    elif color == 'blue':
                        font.color.rgb = RGBColor(0, 0, 225)
                        num_list[2] += 1
                    elif color == 'red':
                        font.color.rgb = RGBColor(255, 0, 0)
                        num_list[3] += 1
                    colored = True
                    break
            if not colored:
                p.add_run(character)
        rows += 1
    # p = document.add_paragraph()
    # p.add_run(filename)

def parse_html(filename):
    with open(filename) as f:
        text = f.read()
    soup = BeautifulSoup(text, 'lxml')
    res_list = []
    font_list = []
    row = 0
    for item in soup.find_all('p'):
        my_text = item.text
        temp_text = unicode(item)[3:-4]
        index = 0
        start = 0
        cur_len = -1
        while start < len(temp_text) and cur_len <= len(temp_text):
            col = temp_text[index:].find("<font")
            if col != -1:
                for i in range(start, col + index):
                    cur_len += 1

                index = index + temp_text[index:].find(">")
                color = ""
                for color_item in colorlist:
                    if temp_text[start:index].find(color_item) != -1:
                        color = color_item
                        break
                font_list.append((row, cur_len + 1, color))
                cur_len += 1
            else:
                for i in range(start, len(temp_text)):
                    cur_len += 1

            index = temp_text[index:].find("</font>") + index + 7
            start = index
        row += 1
        res_list.append(my_text)
    return res_list, font_list

if __name__ == "__main__":
    # document = Document()
    # for i in range(170, 180):
    #     filename = str(i) + ".html"
    #     filepath = "html_result/"+ filename
    #     target_file_name = filename[:filename.find('.')]
    #     res_list, font_list = parse_html(filepath)
    #
    #     document.styles['Normal'].font.name = u'宋体'
    #     document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    #     num_list = [0, 0, 0, 0]
    #     html2word(document, res_list, font_list, num_list, "cn", filename)
    # document.save('doc_result/test175.docx')

    document = Document()
    document.styles['Normal'].font.name = u'宋体'
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    num_list = [0, 0, 0, 0]
    for filename in os.listdir("html_result"):
        print(filename)
        filepath = 'html_result/' + filename
        target_file_name = filename[:filename.find('.')]

        res_list, font_list = parse_html(filepath)
        html2word(document, res_list, font_list, num_list, "cn", filename)
    print num_list
    document.save('doc_result/test_9_24.docx')

